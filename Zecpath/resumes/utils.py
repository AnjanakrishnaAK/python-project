import os
import re
import pdfplumber
import docx


class ResumeTextExtractor:
    @staticmethod
    def extract_pdf(file_path):
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:

                page_text = page.extract_text()

                if page_text:
                    text += page_text + "\n"
        return text
    @staticmethod
    def extract_docx(file_path):
        text = ""
        document = docx.Document(file_path)
        for para in document.paragraphs:
            text += para.text + "\n"
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
        return text
    @staticmethod
    def clean(text):
        if not text:
            return ""
        text = re.sub(r'[•●►▪■★]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r'[^\w\s@.+\-(),:/]', ' ', text)
        return text.strip()
    @classmethod
    def extract(cls, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            raw = cls.extract_pdf(file_path)
        elif ext == ".docx":
            raw = cls.extract_docx(file_path)
        else:
            raise Exception("Unsupported file")
        return cls.clean(raw)


SKILL_LIST = [
    "python",
    "django",
    "java",
    "sql",
    "react",
    "javascript",
    "html",
    "css",
]


class ResumeParser:
    @staticmethod
    def extract_email(text):

        match = re.search(r'\S+@\S+', text)

        return match.group() if match else None
    @staticmethod
    def extract_phone(text):

        match = re.search(r'\+?\d[\d\s\-]{8,15}', text)

        return match.group() if match else None
    @staticmethod
    def extract_skills(text):

        found = []

        text_lower = text.lower()

        for skill in SKILL_LIST:

            if skill in text_lower:
                found.append(skill)

        return ", ".join(found)